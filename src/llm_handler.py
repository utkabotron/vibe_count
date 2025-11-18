"""
Модуль интеграции с OpenAI API
"""
import logging
import json
from pathlib import Path
from openai import OpenAI

from config import Config

logger = logging.getLogger(__name__)


class LLMHandler:
    """Класс для работы с OpenAI API"""

    SYSTEM_PROMPT = """Ты — эксперт по извлечению структурированных данных из бухгалтерских документов (счета на оплату, накладные, УПД). Твоя задача — проанализировать предоставленный файл и извлечь данные в строгом формате JSON.

ПРАВИЛА ИЗВЛЕЧЕНИЯ:
1. Если поле отсутствует в документе, установи значение null.
2. Даты приведи к формату "DD.MM.YYYY".
3. Числовые значения очисти от пробелов и символов валют (например, "10 000,00" -> 10000.00).
4. Для табличной части товаров/услуг найди начало таблицы и извлеки все строки.

СТРУКТУРА JSON:
{
  "document_info": {
    "invoice_number": "Номер счета",
    "invoice_date": "Дата выставления (DD.MM.YYYY)"
  },
  "recipient_details": {
    "bank_name": "Банк получателя",
    "inn": "ИНН получателя",
    "kpp": "КПП получателя",
    "payee_name": "Получатель платежа",
    "bic": "БИК",
    "corr_account": "Корреспондентский счет",
    "current_account": "Расчетный счет"
  },
  "buyer_details": {
    "name": "Покупатель",
    "inn": "ИНН покупателя",
    "kpp": "КПП покупателя"
  },
  "logistics": {
    "consignee": "Грузополучатель",
    "consignor": "Грузоотправитель"
  },
  "items": [
    {
      "name": "Наименование",
      "unit": "Ед. изм.",
      "quantity": число,
      "price_unit": число,
      "total_sum": число,
      "vat_rate": число (например, 20 для 20%),
      "vat_amount": число
    }
  ],
  "totals": {
    "total_without_vat": число,
    "total_vat": число,
    "total_amount": число
  }
}

Верни только валидный JSON без markdown-разметки."""

    def __init__(self):
        """Инициализация клиента OpenAI"""
        self.client = OpenAI(api_key=Config.OPENAI_API_KEY)
        logger.info("OpenAI клиент инициализирован")

    def process_file(self, file_path, file_type):
        """
        Обработать файл через OpenAI API

        Args:
            file_path (Path): Путь к файлу
            file_type (str): Тип файла

        Returns:
            dict: Извлеченные данные в формате JSON
        """
        if file_type in ['image', 'pdf']:
            return self._process_image(file_path)
        elif file_type in ['docx', 'xlsx']:
            return self._process_text(file_path)
        else:
            raise ValueError(f"Неподдерживаемый тип файла: {file_type}")

    def _process_image(self, file_path):
        """
        Обработать изображение или PDF через GPT-4 Vision

        Args:
            file_path (Path): Путь к файлу

        Returns:
            dict: Извлеченные данные
        """
        try:
            import base64

            # Читаем файл и конвертируем в base64
            with open(file_path, 'rb') as image_file:
                image_data = base64.b64encode(image_file.read()).decode('utf-8')

            # Определяем MIME тип
            file_ext = file_path.suffix.lower()
            mime_types = {
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.pdf': 'application/pdf'
            }
            mime_type = mime_types.get(file_ext, 'image/jpeg')

            # Вызываем OpenAI API
            response = self.client.chat.completions.create(
                model=Config.OPENAI_MODEL_VISION,
                messages=[
                    {
                        "role": "system",
                        "content": self.SYSTEM_PROMPT
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Извлеки данные из этого документа:"
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{image_data}"
                                }
                            }
                        ]
                    }
                ],
                response_format={"type": "json_object"},
                max_tokens=4096,
                temperature=0
            )

            # Парсим ответ
            content = response.choices[0].message.content
            data = json.loads(content)

            logger.info(f"Данные успешно извлечены из изображения")
            return data

        except Exception as e:
            logger.error(f"Ошибка при обработке изображения: {e}")
            raise

    def _process_text(self, file_path):
        """
        Обработать текстовый документ через GPT-4 Turbo

        Args:
            file_path (Path): Путь к файлу

        Returns:
            dict: Извлеченные данные
        """
        try:
            # Определяем тип файла и извлекаем текст
            file_ext = file_path.suffix.lower()

            if file_ext in ['.doc', '.docx']:
                text = self._extract_text_from_docx(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                text = self._extract_text_from_xlsx(file_path)
            else:
                raise ValueError(f"Неподдерживаемый формат текстового файла: {file_ext}")

            # Вызываем OpenAI API
            response = self.client.chat.completions.create(
                model=Config.OPENAI_MODEL_TEXT,
                messages=[
                    {
                        "role": "system",
                        "content": self.SYSTEM_PROMPT
                    },
                    {
                        "role": "user",
                        "content": f"Извлеки данные из этого текстового документа:\n\n{text}"
                    }
                ],
                response_format={"type": "json_object"},
                max_tokens=4096,
                temperature=0
            )

            # Парсим ответ
            content = response.choices[0].message.content
            data = json.loads(content)

            logger.info(f"Данные успешно извлечены из текстового документа")
            return data

        except Exception as e:
            logger.error(f"Ошибка при обработке текстового документа: {e}")
            raise

    def _extract_text_from_docx(self, file_path):
        """Извлечь текст из DOCX файла"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Извлекаем текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = '\n'.join(text_parts)
            logger.info(f"Извлечено {len(text)} символов из DOCX")
            return text

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOCX: {e}")
            raise

    def _extract_text_from_xlsx(self, file_path):
        """Извлечь текст из XLSX файла"""
        try:
            import pandas as pd

            # Читаем все листы Excel файла
            excel_file = pd.ExcelFile(file_path)
            text_parts = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Добавляем название листа
                text_parts.append(f"=== Лист: {sheet_name} ===")

                # Конвертируем DataFrame в текст
                # Заменяем NaN на пустые строки
                df = df.fillna('')

                # Форматируем как таблицу
                for _, row in df.iterrows():
                    row_text = ' | '.join(str(val) for val in row.values if str(val).strip())
                    if row_text.strip():
                        text_parts.append(row_text)

            text = '\n'.join(text_parts)
            logger.info(f"Извлечено {len(text)} символов из XLSX")
            return text

        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из XLSX: {e}")
            raise
